# Copyright (c) 2026 Marco De Roni. All rights reserved.
# Licensed under the MIT License — see LICENSE file for details.

import os
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- COLORS ---
GREEN_FILL  = PatternFill("solid", fgColor="C6EFCE")
YELLOW_FILL = PatternFill("solid", fgColor="FFEB9C")
RED_FILL    = PatternFill("solid", fgColor="FFC7CE")
HEADER_FILL = PatternFill("solid", fgColor="1F4E79")

HEADER_FONT  = Font(bold=True, color="FFFFFF", size=11)
BOLD_FONT    = Font(bold=True, size=10)
NORMAL_FONT  = Font(size=10)

THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin")
)


def auto_width(ws):
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max_len + 4, 60)


def style_header_row(ws, row_num: int, num_cols: int):
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=row_num, column=col)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = THIN_BORDER


def pct_fill(pct: float) -> PatternFill:
    if pct >= 80:
        return GREEN_FILL
    elif pct >= 40:
        return YELLOW_FILL
    else:
        return RED_FILL


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


# ─────────────────────────────────────────────
# EXCEL REPORT
# ─────────────────────────────────────────────

def generate_excel(
    keyword_results: dict,
    clause_results: dict,
    metadata_results: dict,
    comparison_results: dict,
    output_dir: str
) -> str:

    wb = Workbook()
    wb.remove(wb.active)

    # ── Sheet 1: Summary ──
    ws = wb.create_sheet("Summary")
    ws.append(["CONTRACT BULK ANALYZER — SUMMARY REPORT"])
    ws["A1"].font = Font(bold=True, size=14, color="1F4E79")
    ws.append([f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}"])
    ws.append([])
    ws.append(["Total contracts analysed", len(metadata_results)])
    ws.append(["Keywords searched", len(keyword_results)])
    ws.append(["Clauses checked", len(clause_results)])
    auto_width(ws)

    # ── Sheet 2: Keyword Frequency ──
    ws2 = wb.create_sheet("Keyword Frequency")
    headers = ["Keyword", "Contracts Found", "Total Contracts", "% Presence", "Confidence"]
    ws2.append(headers)
    style_header_row(ws2, 1, len(headers))

    for kw, data in keyword_results.items():
        row = [kw, data["contracts_found"], data["contracts_total"],
               f"{data['percentage']}%", data.get("confidence", "")]
        ws2.append(row)
        r = ws2.max_row
        ws2.cell(r, 4).fill = pct_fill(data["percentage"])
        ws2.cell(r, 4).font = BOLD_FONT
        for col in range(1, 6):
            ws2.cell(r, col).border = THIN_BORDER
    auto_width(ws2)

    # ── Sheet 3: Clause Presence ──
    ws3 = wb.create_sheet("Clause Presence")
    headers3 = ["Clause", "Present", "Absent", "Total", "% Presence"]
    ws3.append(headers3)
    style_header_row(ws3, 1, len(headers3))

    for name, data in clause_results.items():
        row = [name, data["present_count"], data["absent_count"],
               data["total"], f"{data['presence_pct']}%"]
        ws3.append(row)
        r = ws3.max_row
        ws3.cell(r, 5).fill = pct_fill(data["presence_pct"])
        ws3.cell(r, 5).font = BOLD_FONT
        for col in range(1, 6):
            ws3.cell(r, col).border = THIN_BORDER
    auto_width(ws3)

    # ── Sheet 4: Clause Matrix ──
    ws4 = wb.create_sheet("Clause Matrix")
    clause_names = list(clause_results.keys())
    filenames = list(metadata_results.keys())

    headers4 = ["Contract"] + clause_names
    ws4.append(headers4)
    style_header_row(ws4, 1, len(headers4))

    for filename in filenames:
        row = [filename]
        for clause_name in clause_names:
            present_in = clause_results[clause_name]["present_in"]
            row.append("✓" if filename in present_in else "✗")
        ws4.append(row)
        r = ws4.max_row
        for col in range(2, len(clause_names) + 2):
            cell = ws4.cell(r, col)
            cell.fill = GREEN_FILL if cell.value == "✓" else RED_FILL
            cell.font = BOLD_FONT
            cell.alignment = Alignment(horizontal="center")
            cell.border = THIN_BORDER
        ws4.cell(r, 1).border = THIN_BORDER
    auto_width(ws4)

    # ── Sheet 5: Metadata ──
    ws5 = wb.create_sheet("Metadata")
    meta_fields = ["parties", "effective_date", "governing_law",
                   "jurisdiction", "notice_period", "duration", "auto_renewal"]
    headers5 = ["Contract"] + [f.replace("_", " ").title() for f in meta_fields]
    ws5.append(headers5)
    style_header_row(ws5, 1, len(headers5))

    for filename, meta in metadata_results.items():
        row = [filename] + [meta.get(f, "Not detected") for f in meta_fields]
        ws5.append(row)
        r = ws5.max_row
        for col in range(1, len(headers5) + 1):
            ws5.cell(r, col).border = THIN_BORDER
            ws5.cell(r, col).font = NORMAL_FONT
    auto_width(ws5)

    # ── Sheet 6: Group Comparison ──
    if comparison_results and comparison_results.get("comparison"):
        ws6 = wb.create_sheet("Group Comparison")
        comparison = comparison_results["comparison"]
        group_names = list(comparison_results["grouped"].keys())

        headers6 = ["Clause"] + [f"{g} (%)" for g in group_names]
        ws6.append(headers6)
        style_header_row(ws6, 1, len(headers6))

        for clause_name, group_data in comparison.items():
            row = [clause_name]
            for group in group_names:
                pct = group_data.get(group, {}).get("percentage", 0)
                row.append(f"{pct}%")
            ws6.append(row)
            r = ws6.max_row
            for col in range(2, len(group_names) + 2):
                cell = ws6.cell(r, col)
                try:
                    val = float(str(cell.value).replace("%", ""))
                    cell.fill = pct_fill(val)
                except Exception:
                    pass
                cell.font = BOLD_FONT
                cell.border = THIN_BORDER
            ws6.cell(r, 1).border = THIN_BORDER
        auto_width(ws6)

    # ── Sheet 7: Clause Extracts ──
    ws7 = wb.create_sheet("Clause Extracts")
    headers7 = ["Clause", "Contract", "Extracted Text"]
    ws7.append(headers7)
    style_header_row(ws7, 1, len(headers7))

    for clause_name, data in clause_results.items():
        extracts = data.get("extracts", {})
        for filename, extract in extracts.items():
            ws7.append([clause_name, filename, extract])
            r = ws7.max_row
            for col in range(1, 4):
                ws7.cell(r, col).border = THIN_BORDER
                ws7.cell(r, col).font = NORMAL_FONT
            ws7.cell(r, 3).alignment = Alignment(wrap_text=True)

    ws7.column_dimensions["A"].width = 25
    ws7.column_dimensions["B"].width = 40
    ws7.column_dimensions["C"].width = 80

    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"bulk_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx")
    wb.save(path)
    return path


# ─────────────────────────────────────────────
# WORD REPORT
# ─────────────────────────────────────────────

def generate_word(
    keyword_results: dict,
    clause_results: dict,
    metadata_results: dict,
    comparison_results: dict,
    output_dir: str,
    pii_summary: dict = None
) -> str:

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title = doc.add_heading("CONTRACT BULK ANALYSIS REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}\n"
        f"Contracts analysed: {len(metadata_results)}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- PII Redaction Summary ---
    if pii_summary:
        add_heading(doc, "Privacy & PII Redaction Summary", level=1)
        total = pii_summary.get("total_entities", 0)
        contracts_count = pii_summary.get("contracts_count", 0)
        breakdown = pii_summary.get("breakdown", {})

        doc.add_paragraph(
            f"Before analysis, {total} sensitive entities were automatically redacted "
            f"across {contracts_count} contracts and restored in this report."
        )

        if breakdown:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Entity Type"
            table.rows[0].cells[1].text = "Count"
            for entity_type, count in sorted(breakdown.items()):
                row = table.add_row().cells
                row[0].text = entity_type
                row[1].text = str(count)
        doc.add_paragraph()

    # --- Keyword Frequency ---
    add_heading(doc, "1. Keyword Frequency", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Keyword", "Found In", "Total", "% Presence", "Confidence"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for kw, data in keyword_results.items():
        row = table.add_row().cells
        row[0].text = kw
        row[1].text = str(data["contracts_found"])
        row[2].text = str(data["contracts_total"])
        row[3].text = f"{data['percentage']}%"
        row[4].text = data.get("confidence", "")

    doc.add_paragraph()

    # --- Clause Presence ---
    add_heading(doc, "2. Clause Presence", level=1)
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = "Table Grid"
    for i, h in enumerate(["Clause", "Present", "Absent", "Total", "% Presence"]):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for name, data in clause_results.items():
        row = table2.add_row().cells
        row[0].text = name
        row[1].text = str(data["present_count"])
        row[2].text = str(data["absent_count"])
        row[3].text = str(data["total"])
        row[4].text = f"{data['presence_pct']}%"

    doc.add_paragraph()

    # --- Group Comparison Divergences ---
    if comparison_results and comparison_results.get("divergences"):
        add_heading(doc, "3. Notable Divergences Between Contract Groups", level=1)
        for d in comparison_results["divergences"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(
                f"{d['clause']}: {d['max_group']} {d['max_pct']}% vs "
                f"{d['min_group']} {d['min_pct']}% (gap: {d['gap']}%)"
            )
            run.font.size = Pt(10)
        doc.add_paragraph()

    # --- Metadata Summary ---
    add_heading(doc, "4. Metadata Summary", level=1)
    meta_fields = ["parties", "effective_date", "governing_law",
                   "jurisdiction", "notice_period", "duration", "auto_renewal"]
    table3 = doc.add_table(rows=1, cols=len(meta_fields) + 1)
    table3.style = "Table Grid"
    headers = ["Contract"] + [f.replace("_", " ").title() for f in meta_fields]
    for i, h in enumerate(headers):
        cell = table3.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)

    for filename, meta in metadata_results.items():
        row = table3.add_row().cells
        row[0].text = filename
        for i, f in enumerate(meta_fields):
            row[i + 1].text = meta.get(f, "—")
            row[i + 1].paragraphs[0].runs[0].font.size = Pt(8)

    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"bulk_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
    doc.save(path)
    return path